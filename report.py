from docx import Document

# Recreate a new Document for the brief version after environment reset
doc = Document()

# Title Page
doc.add_heading("Water Distribution Time Series Data Analytics", level=1)
doc.add_paragraph("A B.Tech Project Report\nSubmitted by\n\nYogesh Singh\n\nUnder the Supervision\nOf\nDr. Shobhana Singh\n\nDepartment of Mechanical Engineering\nIndian Institute of Technology Jodhpur\nNovember, 2024")

doc.add_page_break()

# Table of Contents
doc.add_heading("CONTENTS", level=1)
brief_contents = [
    "1. Acknowledgements",
    "2. Abstract",
    "3. Introduction\n\t1.1 Objective\n\t1.2 Background and Motivation",
    "4. Methodology\n\t2.1 Data Collection and Preprocessing\n\t2.2 Pattern Recognition and Feature Extraction",
    "5. Results and Discussion",
    "6. Conclusion",
    "7. References"
]
for item in brief_contents:
    doc.add_paragraph(item)

doc.add_page_break()

# Acknowledgements
doc.add_heading("Acknowledgements", level=1)
doc.add_paragraph("I would like to extend my sincere gratitude to Dr. Shobhana Singh for her invaluable guidance throughout this project. I am grateful to my family, friends, and colleagues for their support.")

doc.add_page_break()

# Abstract
doc.add_heading("Abstract", level=1)
doc.add_paragraph("This project addresses the challenge of detecting pipe bursts in water distribution networks through real-time data analytics. By leveraging time series data on flow, pressure, and velocity, we aim to develop a model that identifies patterns indicating burst risks. Techniques in pattern recognition and anomaly detection are central to this work, providing a pathway for early alerts in smart water systems.")

doc.add_page_break()

# 1. Introduction
doc.add_heading("1. Introduction", level=1)
doc.add_heading("1.1 Objective", level=2)
doc.add_paragraph("This project’s primary objective is to create a predictive framework that detects burst risk in water pipelines by identifying patterns in operational data. Key goals include:\n"
                  "- Recognizing data patterns linked to burst events.\n"
                  "- Developing a model capable of real-time anomaly detection.\n"
                  "- Validating model accuracy and robustness across different pipeline sections.")

doc.add_heading("1.2 Background and Motivation", level=2)
doc.add_paragraph("Water distribution systems face challenges from pipe deterioration and fluctuating demands, making bursts both costly and disruptive. By applying data analytics and machine learning to real-time monitoring, we aim to improve response times and optimize maintenance practices in these networks.")

doc.add_page_break()

# 2. Methodology
doc.add_heading("2. Methodology", level=1)
doc.add_heading("2.1 Data Collection and Preprocessing", level=2)
doc.add_paragraph("Data collected from various nodes included flow, velocity, and pressure. Key preprocessing steps involved removing outliers, normalizing data, and handling missing values, ensuring readiness for pattern analysis.")

doc.add_heading("2.2 Pattern Recognition and Feature Extraction", level=2)
doc.add_paragraph("Pattern recognition methods were used to identify typical and atypical behaviors in the data, particularly changes indicating possible bursts. Statistical and machine learning techniques were applied to extract features such as peak flow rate and rolling averages, which are sensitive to anomalies.")

doc.add_page_break()

# 3. Results and Discussion
doc.add_heading("3. Results and Discussion", level=1)
doc.add_paragraph("Our model successfully identified bursts by recognizing patterns of abnormal pressure drops and flow spikes. Precision and recall metrics showed high accuracy, highlighting the model’s ability to predict potential burst events with minimal false positives.")

doc.add_page_break()

# 4. Conclusion
doc.add_heading("4. Conclusion", level=1)
doc.add_paragraph("The project demonstrates the viability of using time series data analytics for burst prediction in water distribution systems. The model developed not only detects existing bursts but also aids in preemptive maintenance planning. Future improvements could include incorporating environmental factors to enhance model precision.")

doc.add_page_break()

# References
doc.add_heading("References", level=1)
doc.add_paragraph("1. H. Meena, P. George, and S. Singh, \"Demonstration of real-time monitoring in smart water distribution networks for burst detection,\" Journal of Water Supply: Research and Technology - AQUA, 72(11), 2152 (2024).")

# Save the brief document
brief_doc_path = "BTech_Project_Report_Yogesh_Singh.docx"
doc.save(brief_doc_path)

brief_doc_path
